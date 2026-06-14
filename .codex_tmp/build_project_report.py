from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"C:\Users\wlgns\Documents\Github\3DGP_Assignment4\4차과제보고서.docx"
FONT = "Malgun Gothic"


def normalize_tone(text):
    replacements = [
        ("했습니다", "하였습니다"),
        ("했다", "하였습니다"),
        ("였다", "였습니다"),
        ("이다", "입니다"),
        ("된다", "됩니다"),
        ("한다", "합니다"),
        ("갖는다", "갖습니다"),
        ("넘어간다", "넘어갑니다"),
        ("종료된다", "종료됩니다"),
        ("있다", "있습니다"),
        ("없다", "없습니다"),
        ("않는다", "않습니다"),
        ("않았다", "않았습니다"),
        ("찾는다", "찾습니다"),
        ("어렵다", "어렵습니다"),
        ("줄였다", "줄였습니다"),
    ]
    for src, dst in replacements:
        text = text.replace(src, dst)
    return text


def set_run_font(run, size=10.5, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_paragraph(doc, text="", size=10.5, bold=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, before=before, after=after)
    text = normalize_tone(text)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_blank(doc, count=1):
    for _ in range(count):
        add_paragraph(doc, "", after=0)


def add_heading(doc, number, title):
    p = add_paragraph(doc, f"{number}. {title}", size=11.5, bold=True, before=10, after=6)
    return p


def add_item(doc, marker, text):
    return add_paragraph(doc, f"{marker}) {text}", size=10.5, after=4)


def add_bullet(doc, text):
    return add_paragraph(doc, f"· {text}", size=10.5, after=2)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="888888", size="12"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=160, start=160, bottom=160, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_screenshot_placeholder(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.2)
    cell = table.cell(0, 0)
    cell.width = Cm(15.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell, "808080", "12")
    set_cell_margins(cell, top=240, start=240, bottom=240, end=240)
    shade_cell(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, after=0)
    run = p.add_run("\n\n\n\n\n실행 결과 스크린샷 삽입 공간\n\n\n\n\n")
    set_run_font(run, size=11, bold=False)
    run.font.color.rgb = RGBColor(90, 90, 90)
    add_paragraph(doc, "그림 1. 실행 화면 스크린샷 삽입 위치", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

styles = doc.styles
styles["Normal"].font.name = FONT
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
styles["Normal"].font.size = Pt(10.5)

add_blank(doc, 8)
add_paragraph(doc, "3D 게임 프로그래밍 과제 4 설명 문서", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
add_blank(doc, 3)
add_paragraph(doc, "2022184025 윤지훈", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

add_paragraph(doc, "목차", size=11.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
toc_items = [
    "1. 과제 목표",
    "2. 가정",
    "3. 조작법",
    "4. 실행 결과",
    "5. 구현 내용",
    "6. 알고리즘",
]
for item in toc_items:
    add_paragraph(doc, item, size=10.5, after=4)
doc.add_page_break()

add_heading(doc, 1, "과제 목표")
add_paragraph(
    doc,
    "Direct3D 12 기반의 3D 전투 게임 프로그램을 제작한다. 기존 헬리콥터 슈팅 스테이지를 "
    "기반으로 하여, 지형 위에서 진행되는 탱크 전투 스테이지와 헬리콥터 대 탱크 전투 "
    "스테이지를 추가했다. 프로그램은 타이틀 화면에서 시작해 스테이지 선택 화면을 거쳐 "
    "LEVEL-1, LEVEL-2, LEVEL-3 로 진입하는 구조를 갖는다.",
)
add_paragraph(
    doc,
    "LEVEL-1 은 Apache 헬리콥터로 지형 위의 적 객체를 조준하고 미사일로 파괴하는 방식이며, "
    "LEVEL-2 는 Abrams 탱크를 조작하여 적 탱크와 교전하는 방식입니다. LEVEL-3 은 다시 "
    "헬리콥터를 조작하되, 적 객체를 단순 박스가 아닌 탱크로 배치하고, 적 탱크가 플레이어 "
    "헬리콥터를 조준하여 포탄을 발사하도록 구성했습니다.",
)
add_paragraph(
    doc,
    "추가로 라이프 바, 탱크 실드, 자동 공격, 락온 표시, 조종석 1인칭 시점, 거의 투명한 "
    "헬리콥터 유리창, 투명 오브젝트 렌더링 순서 처리를 구현하여 하나의 프로젝트 형태로 "
    "동작하도록 하는 것을 목표로 했습니다.",
)

add_heading(doc, 2, "가정")
assumptions = [
    "C++20 / Windows 10 이상 / Direct3D 12 환경에서 실행한다.",
    "월드 좌표계에서 float 1.0f 는 1m 로 간주한다.",
    "지형은 Textures 폴더의 Cheongsando.png 하이트맵을 기준으로 생성한다. 하이트맵이 없을 경우에는 기본 평면 지형을 사용한다.",
    "Apache, AbramsTank, Rock, Rock2 모델은 Models 폴더의 txt 파일을 읽어 파트 단위 메시로 저장한다.",
    "Apache 모델 로딩에 실패하면 박스 조합으로 만든 헬기 모형을 대신 사용한다.",
    "Abrams 탱크 모델은 파트 이름과 위치를 기준으로 차체, 포탑, 포신을 구분한다.",
    "LEVEL-2 에서 플레이어 탱크는 체력 5, 적 탱크는 체력 2 로 시작한다.",
    "LEVEL-3 에서 플레이어 헬리콥터는 체력 3 으로 시작하고, 적 탱크는 체력 2 로 시작한다.",
    "조명은 방향성 광원 하나만 사용하며, 앰비언트 / 디퓨즈 / 스페큘러 항만 구현한다.",
    "투명한 유리창은 일반 오브젝트를 먼저 렌더링한 뒤, 깊이 쓰기를 끈 투명 전용 파이프라인으로 나중에 렌더링한다.",
]
for i, text in enumerate(assumptions):
    add_item(doc, i, text)

add_heading(doc, 3, "조작법")
add_item(doc, 1, "타이틀 화면")
add_paragraph(doc, "화면 중앙의 PLAY 글자를 좌 클릭하면 글자가 폭발하듯 흩어진 뒤 스테이지 선택 화면으로 넘어간다.", after=4)
add_item(doc, 2, "스테이지 선택 화면")
add_paragraph(doc, "TUTORIAL, LEVEL-1, LEVEL-2, LEVEL-3, START, END 항목이 있으며, START 는 LEVEL-1 과 동일하게 동작한다. END 를 누르면 프로그램이 종료된다.", after=4)
add_item(doc, 3, "LEVEL-1 헬리콥터 스테이지")
for text in [
    "W / A / S / D: 전후좌우 이동",
    "마우스 움직임: 헬리콥터의 Yaw / Pitch 조작",
    "Space: 상승",
    "Left Ctrl: 하강",
    "좌 클릭: 현재 조준 방향으로 미사일 발사",
    "우 클릭: 현재 락온 대상 고정 / 한 번 더 누르면 고정 해제",
    "N 또는 목표 지점 도달: LEVEL-2 로 이동",
    "ESC: 스테이지 선택 화면으로 복귀",
]:
    add_bullet(doc, text)
add_item(doc, 4, "LEVEL-2 탱크 스테이지")
for text in [
    "방향키: 플레이어 탱크 이동",
    "마우스 움직임: 카메라 방향과 포탑 조준 방향 조작",
    "좌 클릭 또는 Space: 포탄 발사",
    "우 클릭: 현재 포신 방향에 있는 적 탱크 선택",
    "A: 자동 공격 켜기 / 끄기",
    "S: 실드 켜기 / 끄기",
    "ESC: 스테이지 선택 화면으로 복귀",
]:
    add_bullet(doc, text)
add_item(doc, 5, "LEVEL-3 헬리콥터 대 탱크 스테이지")
for text in [
    "W / A / S / D: 전후좌우 이동",
    "마우스 움직임: 헬리콥터의 Yaw / Pitch 조작",
    "Space / Left Ctrl: 상승 / 하강",
    "좌 클릭: 미사일 발사",
    "우 클릭: 현재 락온 탱크 고정 / 고정 해제",
    "V: 1인칭 / 3인칭 시점 전환",
    "1: 조종석 1인칭 시점으로 전환",
    "3: 3인칭 추적 시점으로 전환",
    "ESC: 스테이지 선택 화면으로 복귀",
]:
    add_bullet(doc, text)

add_heading(doc, 4, "실행 결과")
add_paragraph(
    doc,
    "프로그램 실행 시 타이틀 화면과 PLAY 버튼이 표시됩니다. PLAY 를 누르면 3D 텍스트가 "
    "폭발하듯 흩어지고, 일정 시간이 지난 뒤 메뉴 화면으로 넘어갑니다.",
)
add_paragraph(
    doc,
    "LEVEL-1 에 진입하면 하이트맵 지형 위에 Apache 헬리콥터와 적 박스들이 배치됩니다. "
    "좌 클릭으로 미사일을 발사하면 헬리콥터 전방으로 날아가고, 락온 대상이 있으면 일정 "
    "시간 후 목표 방향으로 회전하며 추적합니다. 미사일은 비행 중 트레일을 남기고, 지형이나 "
    "적에 닿으면 폭발 파티클을 생성합니다.",
)
add_paragraph(
    doc,
    "LEVEL-2 에서는 플레이어 탱크와 적 탱크가 지형 위에 배치됩니다. 플레이어는 방향키로 "
    "탱크를 이동시키고, 마우스로 포탑과 포신을 조준하며 포탄을 발사합니다. 적 탱크는 "
    "플레이어를 향해 포탑을 회전시키고, 조준이 맞으면 포탄을 발사합니다. 실드가 켜져 있을 "
    "때는 피격 폭발 색상이 다르게 출력되고, 체력은 화면 하단의 라이프 바로 확인할 수 있습니다.",
)
add_paragraph(
    doc,
    "LEVEL-3 에서는 헬리콥터가 적 탱크를 상대합니다. 적 탱크는 헬리콥터를 조준하여 포탄을 "
    "발사하고, 헬리콥터는 미사일과 락온 기능으로 탱크를 공격합니다. 1인칭 시점으로 전환하면 "
    "카메라가 Apache 모델의 조종석 위치로 이동하며, 조종석 유리창은 거의 투명하게 보이도록 "
    "처리했습니다. 유리창 뒤의 적 탱크는 투명 렌더링 파이프라인을 통해 조종석 유리창 너머에서도 "
    "정상적으로 표시됩니다.",
)
add_screenshot_placeholder(doc)

add_heading(doc, 5, "구현 내용")
sections = [
    (
        "1",
        "프레임워크와 자원 로딩",
        [
            "기본 렌더링 흐름은 Direct3D 12 의 디바이스, 스왑체인, 렌더 타깃, 깊이 버퍼, 파이프라인 상태를 생성한 뒤 DrawIndexedInstanced 로 메시를 그리는 방식입니다.",
            "Assets.cpp 에서 Apache.txt, AbramsTank.txt, Rock.txt, Rock2.txt 를 순서대로 읽고, 각 모델을 ModelType 별 핸들로 저장합니다.",
            "모델 파서는 <Frame>, <TransformMatrix>, <Mesh>, <Positions>, <Normals>, <SubMesh>, <AlbedoColor> 태그를 라인 단위로 읽어 파트별 메시를 구성합니다.",
        ],
    ),
    (
        "2",
        "스테이지 구성",
        [
            "GameScene 에 Start, Menu, Tutorial, Level1, Level2, Level3 상태를 두고, 현재 씬에 따라 Update 와 BuildDrawItems 의 동작을 분기했습니다.",
            "LEVEL-1 에서 모든 목표를 파괴하거나 목표 지점에 도달하면 LEVEL-2 로 넘어가도록 했습니다. LEVEL-2 와 LEVEL-3 은 메뉴에서 직접 선택할 수도 있습니다.",
            "각 레벨을 ResetLevel(), ResetLevel2(), ResetLevel3() 으로 초기화하여 플레이어 위치, 적 배치, 탄환, 폭발, 트레일, 락온 상태를 한 번에 정리합니다.",
        ],
    ),
    (
        "3",
        "탱크 모델과 조작",
        [
            "AbramsTank 모델은 파트 이름을 검사하여 포탑과 포신에 해당하는 파트를 구분했습니다. 포탑은 Yaw 회전을 적용하고, 포신은 별도의 Pitch 회전을 적용합니다.",
            "LEVEL-2 의 플레이어 탱크는 방향키 입력을 카메라 기준 forward/right 벡터로 바꾸어 이동합니다. 이동 방향이 생기면 차체 Yaw 를 목표 방향으로 천천히 돌립니다.",
            "탱크는 지형 위 네 지점의 높이를 샘플링하여 차체의 Pitch 와 Roll 을 계산합니다. 이를 통해 경사진 지형 위에서도 모델이 어느 정도 지형을 따라가도록 했습니다.",
        ],
    ),
    (
        "4",
        "적 탱크와 포탄",
        [
            "적 탱크는 매 프레임 플레이어 위치를 향해 포탑과 포신을 회전시킵니다. 조준 방향과 목표 방향의 내적이 일정 값 이상이면 포탄을 발사합니다.",
            "포탄은 Bullet 클래스를 재사용하되 Owner 값을 통해 플레이어 포탄과 적 포탄을 구분했습니다.",
            "포탄과 미사일은 이전 위치와 현재 위치 사이를 광선으로 보고 지형 충돌을 검사하여, 빠르게 움직이는 탄환이 지형을 통과하는 문제를 줄였습니다.",
        ],
    ),
    (
        "5",
        "헬리콥터와 1인칭 조종석",
        [
            "Apache 모델은 PlayerModelWorldMatrix() 하나로 위치, 회전, 스케일을 제어합니다. 주 로터와 꼬리 로터 파트는 중심점 기준으로 따로 회전시켜 움직임을 표현했습니다.",
            "LEVEL-3 에서는 V, 1, 3 키로 헬리콥터 시점을 전환할 수 있습니다. 1인칭 시점에서는 Apache 모델의 로컬 조종석 좌표를 월드 좌표로 변환하여 카메라 위치로 사용합니다.",
            "Apache 모델의 glass 파트는 별도로 찾아 색상 alpha 값을 낮게 적용했습니다. 이를 통해 조종석 내부에서 보았을 때 유리창이 거의 투명하게 보이도록 했습니다.",
        ],
    ),
    (
        "6",
        "조준, 락온, HUD",
        [
            "헬리콥터 조준은 총구 위치에서 전방 방향으로 광선을 쏘고, 탱크 조준은 포구 위치와 포신 방향을 사용합니다.",
            "LEVEL-2 에서는 포신 방향으로 RaycastTankBoundingBox() 를 실행해 우 클릭 시 선택 가능한 탱크를 찾습니다.",
            "LEVEL-3 에서는 화면 중앙에 가까운 탱크를 자동 락온 후보로 고르고, 우 클릭으로 해당 대상을 고정할 수 있게 했습니다.",
            "라이프 바는 NDC 공간에 작은 박스를 여러 개 배치하는 방식으로 구현했습니다. 현재 체력과 최대 체력에 따라 채워지는 칸 수가 달라집니다.",
        ],
    ),
    (
        "7",
        "투명 렌더링 처리",
        [
            "Apache 모델의 glass 파트는 alpha 값이 낮은 투명 DrawItem 으로 분류되도록 하였습니다.",
            "Renderer 에서는 불투명 DrawItem 을 먼저 렌더링하고, alpha 값이 1 보다 작은 DrawItem 은 투명 전용 파이프라인으로 별도 렌더링합니다.",
            "투명 전용 파이프라인은 깊이 검사를 유지하면서 DepthWriteMask 를 ZERO 로 설정합니다. 또한 투명 DrawItem 은 카메라에서 먼 순서로 정렬하여 조종석 유리창 너머의 적 탱크가 계속 보이도록 구성하였습니다.",
        ],
    ),
    (
        "8",
        "파티클과 트레일",
        [
            "폭발은 여러 개의 작은 큐브가 중심에서 바깥으로 퍼지는 방식으로 구현했습니다. seed 값을 사용해 매번 일정한 방향 분포가 나오도록 했습니다.",
            "미사일과 포탄 뒤쪽에는 MissileTrailParticle 을 생성합니다. 트레일은 고정 크기 배열을 순환 사용하여 계속 새 객체를 할당하지 않도록 했습니다.",
        ],
    ),
]
for marker, title, paragraphs in sections:
    add_item(doc, marker, title)
    for text in paragraphs:
        add_paragraph(doc, text, after=4)

add_heading(doc, 6, "알고리즘")
algo_sections = [
    (
        "1",
        "하이트맵 높이 보간",
        [
            "월드 좌표 x/z 를 지형 셀 크기로 나누어 하이트맵 그리드 좌표로 변환합니다. 이후 현재 위치가 포함된 칸의 네 정점 높이를 가져와 x 방향으로 한 번, z 방향으로 한 번 선형 보간합니다.",
            "h0 = lerp(h00, h10, tx);",
            "h1 = lerp(h01, h11, tx);",
            "height = lerp(h0, h1, tz);",
        ],
    ),
    (
        "2",
        "탱크 포탑 / 포신 조준",
        [
            "탱크의 포탑은 현재 포탑 방향과 목표 방향 사이의 Yaw 차이를 구한 뒤, 한 프레임에 회전 가능한 최대 각도만큼만 이동합니다.",
            "포신은 포구 위치에서 목표 지점으로 향하는 방향을 다시 계산하고, 현재 포신 방향과 목표 방향 사이의 Pitch 차이를 구해 제한된 속도로 회전시킵니다.",
            "최종 AimDirection 과 목표 방향의 내적이 TankAutoFireAimDot 이상이면 충분히 조준되었다고 판단하고 발사를 허용합니다.",
        ],
    ),
    (
        "3",
        "탱크 선택과 락온",
        [
            "LEVEL-2 에서는 포구 위치와 포신 방향을 이용해 광선을 만들고, 각 탱크의 OBB 형태 바운딩 박스와 교차하는지 검사합니다. 지형이 먼저 맞으면 그 뒤의 탱크는 선택하지 않습니다.",
            "LEVEL-3 에서는 헬리콥터 전방 방향과 탱크 방향의 내적을 계산합니다. 화면 중앙에 가깝고 거리가 너무 멀지 않은 탱크를 후보로 잡고, 우 클릭 시 해당 대상을 고정합니다.",
            "고정된 탱크가 파괴되면 고정 상태를 해제하고 새 후보를 다시 찾도록 했습니다.",
        ],
    ),
    (
        "4",
        "투명 오브젝트 렌더링",
        [
            "불투명 오브젝트는 기존 파이프라인으로 먼저 그려 깊이 버퍼를 완성합니다.",
            "alpha 값이 1 보다 작은 오브젝트는 별도 배열에 저장한 뒤, 카메라와의 거리 제곱이 큰 순서대로 정렬합니다.",
            "그 다음 DepthWriteMask 를 ZERO 로 둔 투명 파이프라인을 사용해 렌더링합니다. 따라서 유리창은 깊이 검사를 통과한 경우에만 보이지만, 뒤쪽 오브젝트를 깊이 버퍼에서 가려버리지는 않습니다.",
        ],
    ),
    (
        "5",
        "탄환 지형 충돌",
        [
            "탄환은 프레임 사이 이동 거리가 크기 때문에 단순히 현재 위치만 검사하면 지형을 뚫고 지나갈 수 있습니다.",
            "따라서 이전 위치에서 현재 위치로 향하는 광선을 만들고, 그 구간을 일정 간격으로 샘플링하면서 지형 높이보다 아래로 내려가는 순간을 찾습니다.",
            "충돌이 확인되면 해당 위치로 탄환을 이동시키고 수명을 끝낸 뒤 폭발 파티클을 생성합니다.",
        ],
    ),
]
for marker, title, paragraphs in algo_sections:
    add_item(doc, marker, title)
    for text in paragraphs:
        add_paragraph(doc, text, after=4)

doc.save(OUTPUT)
print(OUTPUT)
